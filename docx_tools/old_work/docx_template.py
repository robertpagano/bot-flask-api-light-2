# %%
from docxtpl import DocxTemplate
from docx import Document

# %%
document = Document('articles/Article1.docx')
for i in document.paragraphs:
    print(i)
article_1_full =''
for para in document.paragraphs:
    article_1_full+=para.text

# %%

doc = DocxTemplate("templates/template.docx")
context = { 'article_1_full' : article_1_full }
doc.render(context)
doc.save("generated_docs/test1.docx")
# %%

## above doesn't save any formatting, just text. need to figure out how to get
## formatting into the template stuff

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment



## make logic that if a line has capital letter each word or just one line no 
## punctuation, make it bold?

## look at source code for docxcomposer - maybe i can use this to insert whole
## thing into 

# %%
from docxtpl import DocxTemplate
from docx.shared import Inches

doc = DocxTemplate("templates/template.docx")
sd = doc.new_subdoc(docpath = 'articles/Article1.docx')

context = {
    'mysubdoc': sd,
}

doc.render(context)
doc.save('generated_docs/subdoc_test.docx')

# %%
