# %%
from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer

import requests
import unicodedata

import glob
import os

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import inspect
import win32com.client


# below code grabs a list of the filepaths of all word docs used for testing functions

docx_filepaths = []
for file in glob.glob('doc_builder_files/August/Content and Training/*.docx'):
    docx_filepaths.append(file)


def get_summ(doc_path, url='https://lono-app-v3.azurewebsites.net/api/v1/resources/document/summary/dynamic'):
    '''
    this function calls our summarizing api endpoint, and returns the text of the summary from a word file
    '''

   
    files = {'data': ('file.docx', open(doc_path,'rb'))}
    r = requests.post(url, files=files)

    return unicodedata.normalize("NFKD", r.content.decode('utf-8'))


def create_doc(doc_path, make_summary = False):
    '''
    this function gets called twice in "make_masters" function, once with make_summary == True and once with make_summary == False

    when make_summary == False, it creates a docx object for the actual article contents, and makes article titles "headings" so 
    that the table of contents will grab each article

    when make_summary == True, it creates a docx object the the summary of an article, and also makes article titles "headings"    

    '''
    
    doc = Document(doc_path)
    doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    paragraphs = doc.paragraphs
    title = paragraphs[0].text

    title_2 = paragraphs[0]
    title_2.style = doc.styles['Heading 1']

    if make_summary:
        doc = Document()
        doc.add_heading(title)
        doc.add_paragraph(get_summ(doc_path))

    return doc


def make_toc(doc):
    '''
    this function creates a table of contents object within a docx object, which will be called when the master files are created

    it indexes any text with "heading styles"

    if article "sections" are heading 1, and article titles are heading 2, it will take care of all of the proper formatting/indenting
    '''
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

    return doc

def update_toc(file_name):
    '''
    this has to get called on the master files after all content is populated

    when "make_toc" is called, it creates a "table of contents" object that still needs to be manually "updated" 
    by rightclicking on the object in the doc file, and clicking "update". this function simulates a user opening the file 
    and doing this manually

    '''
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    file_path = os.path.join(script_dir, file_name)
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(file_path)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()

def make_master_file(docx_filepaths, filename, summ=True):
    '''
    for now, this takes in a list of filepaths, and creates two master files: one for summaries, and one for the articles

    it uses the above functions to:
    
        - create the summaries for the text in each individual article
        - formats headings and makes a table of contents for each file

    the input moving forward will likely be doc objects - files will be sent to an api endpoint, the api function will convert to docx objects, then sent to this
    this means we'll need to change this around a bit
    '''

    toc = Document()
    paragraph = toc.add_paragraph('TABLE OF CONTENTS')
    toc = make_toc(toc)
    article_list = [toc]

    for article in docx_filepaths:
        article_list.append(create_doc(article, summ))

    master = article_list[0]
    composer = Composer(master)
    for document in article_list[1:]:
        composer.append(document)

    composer.save(filename) 

    update_toc(filename)


def make_masters(docx_filepaths, summ_filename, article_filename):
    '''
    for now, this takes in a list of filepaths, and creates two master files with function "make_master_file": one for summaries (summ=True), and one for the articles (summ=False)

    it uses the above functions to:
    
        - create the summaries for the text in each individual article
        - formats headings and makes a table of contents for each file

    the input moving forward will likely be doc objects - files will be sent to an api endpoint, the api function will convert to docx objects, then sent to this
    this means we'll need to change this around a bit. May also need to return the actual docx files after saving them. Not sure if the API will handle that or not
    '''

    make_master_file(docx_filepaths, summ_filename, summ=True)
    make_master_file(docx_filepaths, article_filename, summ=False)

make_masters(docx_filepaths, "master_summaries.docx", "master_articles.docx")



# %%
def info_from_path(path):
    path = path.split('/')
    section = path[-2]
    month = path[-3]
    file_name = path[-1]
    article_name = file_name[:-5]
    print(section, month, file_name, article_name)
    return section, month, file_name, article_name

# paths = [
#   "/Shared Documents/TestArticles/August/Content and Training/MSX Integration for intelligent guided selling experience.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Learn about FY21 key themes.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Dynamics 365 Fraud Protection licensing scenarios training.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Quick-Start Guide to Building Resiliency with Customers.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Updated D365 Technical Content.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Infrastructure update.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Technical readiness webinar series.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/What’s New for 2020 release wave 1 webinar series.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/New Business Value Insights templates for rapid response.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Three new Microsoft Dynamics 365 webinar series available on Microsoft.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Review the Quick.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/Quick.docx",
#   "/Shared Documents/TestArticles/August/Product and Availability/tester3.docx",
#   "/Shared Documents/TestArticles/August/Content and Training/tester3.docx"
# ]

# for path in paths:
#     info_from_path(path)

# %%
