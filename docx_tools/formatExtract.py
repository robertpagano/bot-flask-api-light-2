#%%
from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer

# %% This grabs the document from file path
document = Document('data/formattingone.docx')
# %% Extract all text from word document, ignores tables and images.
text =''
for para in document.paragraphs:
    text+=para.text
#%% Adds the summary of the text as a new paragraph at the bottom of the text
from textsummarization import *
document.add_paragraph(extract_sum(text, 0.5))
#%% Saves new document with original formatting + summary at the bottom
document.save('newFormat.docx')
# %% Reads in another word document
document_two = Document('data/formattingtwo.docx')

# %% This cell combines two word documents with formatting into one "combined.docx"
files = ['data/formattingone.docx', 'data/formattingtwo.docx']

master = Document("data/formattingtwo.docx")
composer = Composer(master)
doc1 = Document("data/formattingone.docx")
composer.append(doc1)
composer.save("combined.docx")

# %% Extracts all the tables within a word doc and allows you to add the table to a new word doc
new_doc = Document()
tables = document.tables
tables = tables[0]
tbl = tables._tbl
paragraph = new_doc.add_paragraph()
paragraph._p.addnext(tbl)
new_doc.save('table.docx')
