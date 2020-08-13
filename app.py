from flask import Flask, request, redirect, url_for, flash, jsonify, send_file, render_template, request
# from flask_bootstrap import Bootstrap
from werkzeug.exceptions import Forbidden, HTTPException, NotFound, RequestTimeout, Unauthorized
import numpy as np
import pandas as pd
import pickle as p
import json
import io

from shareplum import Site
from shareplum import Office365
from shareplum.site import Version

from summarization.textsummarization import bert_sum, bert_sum_dynamic
from connector import call_for_content
from docx import Document
from docx.shared import Inches

from docx_tools import word_doc_gen

from linkcheck import flag_private_urls, flag_private_urls_to_dict

app = Flask(__name__)

## Testing out HTML
@app.route('/')
def form():
  return render_template('index.html')

## Handling Errors
@app.errorhandler(NotFound)
def page_not_found_handler(e: HTTPException):
    return render_template('404.html'), 404


@app.errorhandler(Unauthorized)
def unauthorized_handler(e: HTTPException):
    return render_template('401.html'), 401


@app.errorhandler(Forbidden)
def forbidden_handler(e: HTTPException):
    return render_template('403.html'), 403


@app.errorhandler(RequestTimeout)
def request_timeout_handler(e: HTTPException):
    return render_template('408.html'), 408
    
## summarizer using form data as an input, returns text summary
@app.route('/api/v1/resources/text/', methods=['POST'])
def summarize_from_text():
    
    data = request.form["data"]
    summary = bert_sum(data)

    return summary

## summarizer that pulls text from word doc, returns text summary
@app.route('/api/v1/resources/document/summary', methods=['GET', 'POST'])
def summarize_from_file():
    
    f = request.files['data']
    f.save('datafile.docx')
    document = Document('datafile.docx')
    text =''
    for para in document.paragraphs:
        text+=para.text
    
    summary = bert_sum(text)

    return summary

## summarizer that pulls text from word doc, returns text summary with a dynamic length
@app.route('/api/v1/resources/document/summary/dynamic', methods=['GET', 'POST'])
def summarize_from_file_dynamic_length():
    
    f = request.files['data']
    f.save('datafile.docx')
    document = Document('datafile.docx')
    text =''
    for para in document.paragraphs[1:]:
        text+=para.text
    
    summary = bert_sum_dynamic(text)

    return summary

## this function just takes in a word file, and returns the word file
@app.route('/api/v1/resources/document/docx', methods=['GET', 'POST'])
def return_document():
    f = request.files['data']
    f.save('datafile.docx')

    return send_file('datafile.docx', attachment_filename='test.docx')

## this takes in a word file, summarizes text, adds summary to end, and returns document. Used downstream in transform route below
def transform():
    data = 'datafile.docx'
    document = Document(data)
    new = Document()
    text =''
    for para in document.paragraphs[1:]:
        text+=para.text
    new.add_paragraph(bert_sum(text))
    return new

## this uploads the word document to a sharepoint folder
def upload_sp(section):
    data = 'datafile.docx'
    document = Document(data)
    title = document.paragraphs[0].text + ".docx"
    summary_title = document.paragraphs[0].text + "_summary.docx"
    authcookie = Office365('https://thespurgroup.sharepoint.com', username='kevin.lin@thespurgroup.com', password='zcmzlpbxcvtzqwzp').GetCookies()
    site = Site('https://thespurgroup.sharepoint.com/sites/bot_project_test/', version=Version.v2016, authcookie=authcookie)
    folder_month = site.Folder('Shared Documents/Lono2docs/Assets & Templates/')
    month = str(folder_month.get_file('last_call.txt'))[2:-1]
    folder_path = 'Shared Documents/Lono2docs/Newsletter content/' + month + '/' + section + '/'
    folder_main = site.Folder(folder_path)

    with open('datafile.docx', mode='rb') as file:
        fileContent = file.read()
        file.close()
    summary = transform()
    summary.save('summary.docx')
    with open ('summary.docx', mode='rb') as file:
        summaryContent = file.read()
        file.close()
    folder_main.upload_file(fileContent, title)
    folder_main.upload_file(summaryContent, summary_title)

## shows html interface, for now users can submit a document, and it will add a summary to the end
@app.route('/transform', methods=["GET","POST"])
def transform_view():
    f = request.files['data_file']
    section = request.form['section']
    f.save('datafile.docx')
    if not f:
        return "Please upload a word document"
    result = transform()
    result.save('result.docx')

    ## This uploads the file to the correct folder in sharepoint
    upload_sp(section)

    return send_file('result.docx', attachment_filename='new_file.docx')

## this takes in a word file, scrapes for links, checks links, and returns an excel file with the results of link checker
@app.route('/api/v1/resources/document/links/excel', methods=['POST'])
def check_links_to_excel():
    
    data = request.files["link"]
    data.save('linkfile.docx')
    results = flag_private_urls('linkfile.docx')
    results.to_excel('links.xlsx')
    
    return send_file('links.xlsx')

## this takes in a word file, scrapes for links, checks links, and returns a json object of the table of results of link checker
@app.route('/api/v1/resources/document/links/json', methods=['POST'])
def check_links_to_json():
    
    data = request.files["link"]
    data.save('linkfile.docx')
    results = flag_private_urls_to_dict('linkfile.docx')
    
    return jsonify(results)

## this takes in multiple word files, and then will create master summary and article files
## i think we will need to tweak the function and also will need to take in some variables like:
#### month
#### section
#### title

@app.route('/api/v1/resources/document/docbuilder', methods=["POST"])
def build_docs():
    '''
    this takes in two arrays - one for files and one for filepaths. 
    It then creates a dictionary with the following structure:
    
    {article_name: {
        doc: (docx object)
        month: (month)
        section: (section)}
    }
    this will then be used to create the master document files
    '''

    uploaded_files = request.form.getlist("file")
    path_list = request.form.getlist("paths")
    path_list= path_list[0]
    uploaded_files = uploaded_files[0]
    doc_list = []
    for i in uploaded_files:
        doc_list.append(i)
    for document in uploaded_files:
        #make each into a docx object
        source_stream = io.BytesIO(document)
        doc = Document(source_stream)
        doc_list.append(doc)
    final_dict = word_doc_gen.make_doc_dict(path_list, doc_list)
    return str(final_dict)

@app.route('/api/v1/resources/document/links/test', methods=['POST'])
def check_links_to_json_test():
    
    data = request.form["test"]
    source_stream = io.StringIO(data)
    doc = Document(source_stream)
    doc.save('linkfiletest.docx')
    # k = type(data)
    return str(doc)


@app.route("/upload", methods=["POST"])
def upload():
    uploaded_files = request.files.to_dict()
    for file in uploaded_files:
        print(uploaded_files[file].filename)

    return ""

# images = request.files.to_dict() #convert multidict to dict
# for image in images:     #image will be the key 
#     print(images[image])        #this line will print value for the image key
#     file_name = images[image].filename
#     images[image].save(some_destination)

@app.route("/api/v1/resources/document/call", methods=["POST"])
def content_call():
    call_for_content()

if __name__ == '__main__':
    app.run(debug=True)
